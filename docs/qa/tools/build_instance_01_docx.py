from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
SOURCE = ROOT / "docs" / "qa" / "instancia-01-base-funcional.md"
OUTPUT = ROOT / "docs" / "qa" / "grana-v3-qa-instancia-01-base-funcional.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CDD5DF"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, *, font="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_style(style, *, font="Calibri", size=11, color=None, bold=None, before=0, after=6, line=1.25):
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:ascii"), font)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_inline(paragraph, text):
    pieces = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run(run, font="Consolas", size=9.5, color=DARK_BLUE)
        elif piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            set_run(run, bold=True)
        else:
            run = paragraph.add_run(piece)
            set_run(run)


def add_table(doc, rows):
    cols = max(len(row) for row in rows)
    if cols == 2:
        widths = [2700, 6660]
    elif cols == 3:
        widths = [2200, 3600, 3560]
    elif cols == 4:
        widths = [1700, 3000, 1800, 2860]
    elif cols == 5:
        widths = [900, 1950, 1400, 2400, 2710]
    else:
        widths = [TABLE_WIDTH // cols] * cols
        widths[-1] += TABLE_WIDTH - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx in range(cols):
            cell = table.cell(row_idx, col_idx)
            value = row[col_idx] if col_idx < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            run = p.add_run(value)
            set_run(run, size=9.2, bold=(row_idx == 0), color=(NAVY if row_idx == 0 else None))
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        parts = [part.strip() for part in lines[i].strip().strip("|").split("|")]
        if not all(set(part) <= set("-: ") for part in parts):
            rows.append(parts)
        i += 1
    return rows, i


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    set_style(styles["Normal"], after=6, line=1.18)
    set_style(styles["Heading 1"], size=16, color=BLUE, bold=True, before=18, after=10, line=1.0)
    set_style(styles["Heading 2"], size=13, color=BLUE, bold=True, before=14, after=7, line=1.0)
    set_style(styles["Heading 3"], size=11.5, color=DARK_BLUE, bold=True, before=10, after=5, line=1.0)
    for list_style_name in ("List Bullet", "List Number"):
        style = styles[list_style_name]
        set_style(style, after=3, line=1.16)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("GRANA V3  |  QA FUNCIONAL")
    set_run(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Instancia 01  |  Pagina ")
    set_run(run, size=8.5, color=MUTED)
    add_page_field(footer)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("GUIA OPERATIVA DE QA")
    set_run(run, size=10, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Grana V3 - Instancia 01")
    set_run(run, size=25, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Base funcional: auth, onboarding, cuentas y smoke de superficies principales")
    set_run(run, size=12.5, color=MUTED, italic=True)

    metadata = [
        ["Version", "2026-06-01", "Referencia contable", "2026-06-01 (America/Argentina/Buenos_Aires)"],
        ["Alcance", "Web completo", "Cross-platform", "Casos indicados para web y mobile"],
    ]
    add_table(doc, metadata)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    started = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped == "# Grana V3 - Guia QA funcional":
            i += 1
            continue
        if stripped == "## Instancia 01: base funcional":
            started = True
            i += 1
            continue
        if not started:
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(style="Heading 3")
            add_inline(p, stripped[4:])
        elif stripped.startswith("## "):
            p = doc.add_paragraph(style="Heading 1")
            add_inline(p, stripped[3:])
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", stripped))
        elif stripped.endswith(":") and len(stripped) < 80:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, f"**{stripped}**")
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        i += 1

    props = doc.core_properties
    props.title = "Grana V3 - Guia QA funcional - Instancia 01"
    props.subject = "Base funcional: auth, onboarding, cuentas y smoke"
    props.author = "Grana"
    props.keywords = "QA, Grana V3, auth, onboarding, accounts"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()

